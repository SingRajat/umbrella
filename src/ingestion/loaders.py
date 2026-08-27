"""Format-specific document loaders converting raw files into RawDocument structures."""
import os
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional

import docx
import pdfplumber
from pypdf import PdfReader

from src.common.errors import IngestionError
from src.common.logging import get_logger

logger = get_logger("umbrella.ingestion.loaders")


@dataclass
class RawPage:
    """Represents a single page or segment of extracted text."""
    page_num: Optional[int]
    text: str


@dataclass
class RawDocument:
    """Standardized internal representation of an ingested document across all formats."""
    doc_id: str
    filename: str
    source_type: str
    uploaded_at: str
    pages: List[RawPage] = field(default_factory=list)


def load_pdf(file_path: str, doc_id: Optional[str] = None) -> RawDocument:
    """
    Extract text and page metadata from a PDF file.

    Args:
        file_path: Absolute or relative path to the PDF file.
        doc_id: Optional UUID for the document. If None, a new UUID4 is minted.

    Returns:
        RawDocument containing extracted text per page.

    Raises:
        IngestionError: If the PDF is corrupted, unreadable, or contains no text.
    """
    path = Path(file_path)
    if not path.exists():
        raise IngestionError(f"PDF file not found: {file_path}", error_code="FILE_NOT_FOUND", status_code=404)

    minted_doc_id = doc_id or str(uuid.uuid4())
    uploaded_at = datetime.now(timezone.utc).isoformat()
    pages: List[RawPage] = []

    try:
        # Primary loader: pdfplumber (preserves layout and heading awareness)
        with pdfplumber.open(str(path)) as pdf:
            if len(pdf.pages) == 0:
                raise IngestionError("PDF contains 0 pages.", error_code="EMPTY_PDF", status_code=422)

            for idx, page in enumerate(pdf.pages, start=1):
                extracted = page.extract_text() or ""
                pages.append(RawPage(page_num=idx, text=extracted))

    except IngestionError:
        raise
    except Exception as plumber_err:
        logger.warning(f"pdfplumber extraction failed for {path.name}, attempting fallback: {plumber_err}")
        try:
            # Fallback loader: pypdf
            reader = PdfReader(str(path))
            pages = []
            for idx, page in enumerate(reader.pages, start=1):
                extracted = page.extract_text() or ""
                pages.append(RawPage(page_num=idx, text=extracted))
        except Exception as pypdf_err:
            logger.error(f"All PDF extraction methods failed for {path.name}: {pypdf_err}")
            raise IngestionError(
                f"Failed to parse PDF file '{path.name}': corrupted or unreadable format.",
                error_code="CORRUPTED_PDF",
                status_code=422,
            ) from pypdf_err

    # Check total extracted text across all pages
    total_text = "".join(p.text.strip() for p in pages)
    if not total_text:
        raise IngestionError(
            f"PDF file '{path.name}' contains no readable text.",
            error_code="EMPTY_PDF_TEXT",
            status_code=422,
        )

    logger.info(f"Successfully loaded PDF '{path.name}' with {len(pages)} pages (doc_id: {minted_doc_id})")

    return RawDocument(
        doc_id=minted_doc_id,
        filename=path.name,
        source_type="pdf",
        uploaded_at=uploaded_at,
        pages=pages,
    )


def load_docx(file_path: str, doc_id: Optional[str] = None) -> RawDocument:
    """
    Extract text and structure from a DOCX file.

    Args:
        file_path: Absolute or relative path to the DOCX file.
        doc_id: Optional UUID for the document. If None, a new UUID4 is minted.

    Returns:
        RawDocument with extracted paragraphs and tables.

    Raises:
        IngestionError: If the DOCX is corrupted, unreadable, or contains no text.
    """
    path = Path(file_path)
    if not path.exists():
        raise IngestionError(f"DOCX file not found: {file_path}", error_code="FILE_NOT_FOUND", status_code=404)

    minted_doc_id = doc_id or str(uuid.uuid4())
    uploaded_at = datetime.now(timezone.utc).isoformat()

    try:
        doc = docx.Document(str(path))
        extracted_elements = []

        # 1. Extract paragraphs with heading preservation
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if para.style and para.style.name.startswith("Heading"):
                    extracted_elements.append(f"\n## {text}\n")
                else:
                    extracted_elements.append(text)

        # 2. Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    extracted_elements.append(row_text)

    except Exception as exc:
        logger.error(f"Failed to read DOCX file '{path.name}': {exc}")
        raise IngestionError(
            f"Failed to parse DOCX file '{path.name}': corrupted or unreadable format.",
            error_code="CORRUPTED_DOCX",
            status_code=422,
        ) from exc

    full_text = "\n\n".join(extracted_elements).strip()
    if not full_text:
        raise IngestionError(
            f"DOCX file '{path.name}' contains no readable text.",
            error_code="EMPTY_DOCX_TEXT",
            status_code=422,
        )

    logger.info(f"Successfully loaded DOCX '{path.name}' with {len(extracted_elements)} blocks (doc_id: {minted_doc_id})")

    # DOCX files do not have native page boundaries; page_num is None per PRD §3.1
    return RawDocument(
        doc_id=minted_doc_id,
        filename=path.name,
        source_type="docx",
        uploaded_at=uploaded_at,
        pages=[RawPage(page_num=None, text=full_text)],
    )


def load_txt(file_path: str, doc_id: Optional[str] = None) -> RawDocument:
    """
    Extract text from a plain TXT or Markdown (.md) file.

    Args:
        file_path: Absolute or relative path to the TXT or MD file.
        doc_id: Optional UUID for the document. If None, a new UUID4 is minted.

    Returns:
        RawDocument containing file text.

    Raises:
        IngestionError: If the file is not found, unreadable, or contains only whitespace.
    """
    path = Path(file_path)
    if not path.exists():
        raise IngestionError(f"File not found: {file_path}", error_code="FILE_NOT_FOUND", status_code=404)

    minted_doc_id = doc_id or str(uuid.uuid4())
    uploaded_at = datetime.now(timezone.utc).isoformat()
    source_type = "md" if path.suffix.lower() == ".md" else "txt"

    # Attempt UTF-8 with fallback to Latin-1
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
    except UnicodeDecodeError:
        try:
            with open(path, "r", encoding="latin-1") as f:
                content = f.read()
        except Exception as read_err:
            logger.error(f"Failed to decode text file '{path.name}': {read_err}")
            raise IngestionError(
                f"Failed to read file '{path.name}': encoding error.",
                error_code="ENCODING_ERROR",
                status_code=422,
            ) from read_err
    except Exception as exc:
        logger.error(f"Error reading file '{path.name}': {exc}")
        raise IngestionError(
            f"Failed to read file '{path.name}': {exc}",
            error_code="READ_ERROR",
            status_code=422,
        ) from exc

    clean_content = content.strip()
    if not clean_content:
        raise IngestionError(
            f"Text file '{path.name}' contains no readable text.",
            error_code="EMPTY_TXT_TEXT",
            status_code=422,
        )

    logger.info(f"Successfully loaded {source_type.upper()} file '{path.name}' ({len(clean_content)} chars) (doc_id: {minted_doc_id})")

    # TXT / MD files have no native page breaks; single page with page_num: None per PRD §3.1
    return RawDocument(
        doc_id=minted_doc_id,
        filename=path.name,
        source_type=source_type,
        uploaded_at=uploaded_at,
        pages=[RawPage(page_num=None, text=content)],
    )
